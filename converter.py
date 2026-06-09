"""
UNGREED-PDF — Core Conversion Engine
Converts any PDF (text-based or scanned/image) into an editable Word .docx.
Zero external dependencies — Tesseract is bundled, PyMuPDF handles rendering.
"""

from __future__ import annotations

import os
import io
import sys
import fitz  # PyMuPDF
import pytesseract
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from typing import Optional, Callable


# Minimum character count to consider a page as having extractable text
MIN_TEXT_CHARS = 30


def _get_bundle_dir() -> str:
    """Return the directory where bundled files live (PyInstaller or script dir)."""
    if getattr(sys, "frozen", False):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


def _setup_tesseract():
    """Find Tesseract: bundled first, then common Windows paths, then PATH."""
    bundle = _get_bundle_dir()

    # 1. Bundled Tesseract (inside PyInstaller .exe)
    bundled_exe = os.path.join(bundle, "tesseract", "tesseract.exe")
    if os.path.isfile(bundled_exe):
        pytesseract.pytesseract.tesseract_cmd = bundled_exe
        os.environ["TESSDATA_PREFIX"] = os.path.join(bundle, "tesseract", "tessdata")
        return

    # 2. Common Windows install paths
    if sys.platform == "win32":
        for path in [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.join(os.environ.get("LOCALAPPDATA", ""), "Programs", "Tesseract-OCR", "tesseract.exe"),
        ]:
            if os.path.isfile(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return

    # 3. Fall back to system PATH (Linux/macOS or user-installed)


_setup_tesseract()


def _is_text_page(page: fitz.Page) -> bool:
    """Check if a page has enough extractable text to skip OCR."""
    text = page.get_text("text").strip()
    return len(text) >= MIN_TEXT_CHARS


def _extract_text_from_page(page: fitz.Page) -> str:
    """Extract text from a text-based PDF page using PyMuPDF."""
    return page.get_text("text")


def _render_page_to_pil(page: fitz.Page, dpi: int = 300) -> Image.Image:
    """Render a PDF page to a PIL Image using PyMuPDF (no Poppler needed)."""
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat)
    img_data = pix.tobytes("png")
    return Image.open(io.BytesIO(img_data))


def _ocr_page_image(pil_image: Image.Image, lang: str = "eng") -> str:
    """Run Tesseract OCR on a PIL image and return the extracted text."""
    return pytesseract.image_to_string(pil_image, lang=lang)


def _extract_images_from_page(page: fitz.Page) -> list:
    """Extract embedded images from a PDF page as (image_bytes, ext) tuples."""
    images = []
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        try:
            base_image = page.parent.extract_image(xref)
            if base_image:
                images.append((base_image["image"], base_image["ext"]))
        except Exception:
            continue
    return images


def convert_pdf_to_docx(
    pdf_path: str,
    output_path: str | None = None,
    progress_callback: Callable | None = None,
    ocr_lang: str = "eng",
) -> str:
    """
    Convert a PDF file to an editable Word .docx.

    Parameters
    ----------
    pdf_path : str
        Path to the input PDF.
    output_path : str | None
        Path for the output .docx. Defaults to same name/location as PDF.
    progress_callback : callable | None
        Called with (current_page, total_pages) for progress updates.
    ocr_lang : str
        Tesseract language code(s), e.g. "eng", "eng+fra".

    Returns
    -------
    str
        Path to the generated .docx file.
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    if output_path is None:
        base = os.path.splitext(pdf_path)[0]
        output_path = base + ".docx"

    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)

    for page_num in range(total_pages):
        page = pdf_doc[page_num]

        if page_num > 0:
            doc.add_page_break()

        if _is_text_page(page):
            # --- Text-based page: direct extraction ---
            text = _extract_text_from_page(page)
            _add_text_to_doc(doc, text)

            # Also embed any inline images from the page
            images = _extract_images_from_page(page)
            for img_bytes, ext in images:
                _add_image_to_doc(doc, img_bytes, ext)
        else:
            # --- Image / scanned page: render with PyMuPDF → OCR ---
            pil_image = _render_page_to_pil(page, dpi=300)
            ocr_text = _ocr_page_image(pil_image, lang=ocr_lang)
            if ocr_text.strip():
                _add_text_to_doc(doc, ocr_text)
            else:
                # If OCR returns nothing, embed the page as an image
                _add_pil_image_to_doc(doc, pil_image)

        if progress_callback:
            progress_callback(page_num + 1, total_pages)

    pdf_doc.close()
    doc.save(output_path)
    return output_path


def _add_text_to_doc(doc: Document, text: str):
    """Add text to a Word document, preserving paragraph breaks."""
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        stripped = para_text.strip()
        if stripped:
            doc.add_paragraph(stripped)


def _add_image_to_doc(doc: Document, img_bytes: bytes, ext: str):
    """Add an image from bytes to the Word document."""
    try:
        stream = io.BytesIO(img_bytes)
        doc.add_picture(stream, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _add_pil_image_to_doc(doc: Document, pil_image: Image.Image):
    """Add a PIL image to the Word document."""
    try:
        stream = io.BytesIO()
        pil_image.save(stream, format="PNG")
        stream.seek(0)
        doc.add_picture(stream, width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
